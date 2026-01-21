import re
from docx import Document

PLACEHOLDER_PATTERN = r"\{\{.*?\}\}"

def replace_placeholders(doc, mappings):
    
    # -------- BODY CONTENT --------
    for paragraph in doc.paragraphs:
        updated_text = paragraph.text

        for placeholder, value in mappings.items():
            updated_text = updated_text.replace(placeholder, str(value))

        if updated_text != paragraph.text:
            paragraph.clear()
            paragraph.add_run(updated_text)

    # -------- HEADER CONTENT --------
    for section in doc.sections:
        header = section.header

        for paragraph in header.paragraphs:
            for run in paragraph.runs:
                matches = re.findall(PLACEHOLDER_PATTERN, run.text)

                for match in matches:
                    # Replace ONLY if placeholder exists in mapping
                    if match in mappings:
                        run.text = run.text.replace(match, str(mappings[match]))
                   
