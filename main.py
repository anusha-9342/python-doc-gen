from docx import Document
document=Document()
document.add_heading('My Heading') 
document.add_paragraph('This is a sample paragraph in the document.')
document.save('MyDocument.docx')

import pandas as pd
df=pd.read_csv('datasets/sample.csv')
print(df.head())
print(df.columns)

df1=pd.read_excel('datasets/sample.xlsx')
print(df1.head())
print(df1.columns)
